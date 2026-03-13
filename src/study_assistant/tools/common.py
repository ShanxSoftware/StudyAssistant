# src/study_assistant/tools/common.py

from pathlib import Path
import hashlib
import fitz
from enum import Enum
from docx import Document
from study_assistant.config import DB_PATH, MAX_SECTION_TOKENS
from xaihandler.memorystore import MemoryStore
from typing import Optional, List
from pydantic import BaseModel, Field, ConfigDict

def get_doc_key(file_path: Path) -> str: 
    return hashlib.md5(file_path.read_bytes()).hexdigest()
    
class Chunk_Type(Enum):
    PARAGRAPH = "paragraph"
    PAGE = "page" 
    SECTION = "section"
    TABLE = "table"
    FIGURE = "figure"

class Document_Chunk(BaseModel):
    model_config = ConfigDict(use_enum_values=True)
    index: int = Field(..., description="Page Number this chunk starts on")
    page: int = Field(..., description="Number of page, paragraph ")
    chunk_type: Chunk_Type = Field(..., description="Type of chunk")
    text: str = Field(..., description="text value of this chunk")
    tokens_est: int = Field(..., description="Estimated tokens this chunk will cost to process")

class Payload_Metadata(BaseModel): 
    source_path: str = Field(..., description="Path of the file being analysed")
    doc_key: str = Field(..., description="MD5 Document Hash")
    total_chunks: int = Field(..., description="Total number of chunks in this document")
    total_tokens_est: int = Field(..., description="")
    title: str = Field(..., description="Title of the Document")
    chunks: List[Document_Chunk] = Field(..., description="List of chunks for processing")

class Chunked_Document(BaseModel):
    doc_key: str = Field(..., description="MD5 Document Hash")
    metadata: Payload_Metadata = Field(..., description="Payload metadata that describes the document, estimated tokens for processing and the number of chunks")
    chunks: List[Document_Chunk] = Field(..., description="The components of a document broken up into usable chunks")

class Document_Payload(BaseModel): 
    doc_key: Optional[str] = Field(default="", description="MD5 Document Hash")
    chunk_index: Optional[int] = Field(default=0, description="Index of this chunk in list of stored chunks")
    total_chunks: Optional[int] = Field(default=0, description="Total number of chunks")
    progress: Optional[str] = Field(default="", description="index/total_chunks")
    chunk: Optional[Document_Chunk] = Field(default=None, description="the next chunk to be analysed")
    tokens_est: Optional[int] = Field(default=0, description="Estimated Tokens for this document")
    status: str = Field(..., description="Status of this Payload")

def smart_chunk_document(file_path: Path, max_tokens: int = None) -> list[Document_Chunk]:
    """Shared section-aware chunker for both PDF and DOCX."""
    if max_tokens is None: 
        max_tokens = MAX_SECTION_TOKENS

    chunks: list[Document_Chunk] = []
    current_section_text = ""
    current_section_tokens = 0
    current_level = None
    section_index = 0
    suffix = file_path.suffix.lower()
    last_chunk_text = "" # for overlap

    if suffix == ".pdf":
        doc = fitz.open(file_path)
        for page_num, page in enumerate(doc):
            text = page.get_text("text").strip()
            if not text: 
                continue
            # For PDFs we treat each page as potential section boundary (or use font size for headings in future)
            if current_section_tokens > max_tokens or page_num == 0:
                if current_section_text:
                    overlap = last_chunk_text[-150:] if len(last_chunk_text) > 150 else last_chunk_text
                    final_text = (overlap + "\n\n" + current_section_text.strip()).strip()
                    chunks.append(Document_Chunk(
                        index=len(chunks),
                        page=section_index + 1,
                        chunk_type=Chunk_Type.PAGE,
                        text=final_text,
                        tokens_est=current_section_tokens
                    ))
                    section_index += 1
                    last_chunk_text = current_section_text
                current_section_text = text + "\n\n"
                current_section_tokens = int(len(text.split()) * 1.3)
            else: 
                current_section_text += text + "\n\n"
                current_section_tokens += int(len(text.split()) * 1.3)
        doc.close()
    elif suffix == ".docx":
        doc = Document(file_path)
        for para_num, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text: 
                continue
        
            level = None
            style = para.style
            if hasattr(style, 'level'):
                level = style.level # TODO: Verify that this section is organising heading levels and body text into approriate sections. 
            elif para.style.name.startswith("Heading"): # Fallback
                level = int(para.style.name[-1]) if para.style.name[-1].isdigit() else 1

            # Start new section on heading change or when we hit the token limit
            if level is not None or current_section_tokens > max_tokens: 
                if current_section_text:
                    overlap = last_chunk_text[-150:] if len(last_chunk_text) > 150 else last_chunk_text
                    final_text = (overlap + "\n\n" + current_section_text.strip()).strip()
                    #Finalise previous section
                    chunks.append(Document_Chunk(
                        index=len(chunks), 
                        page=section_index + 1, 
                        chunk_type=Chunk_Type.SECTION, 
                        text=final_text,
                        tokens_est=current_section_tokens
                    ))
                    section_index += 1
                    last_chunk_text = current_section_text
                current_section_text = text + "\n\n"
                current_section_tokens = int(len(text.split())*1.3)
                current_level = level
            else: 
                current_section_text += text + "\n\n"
                current_section_tokens += int(len(text.split()) * 1.3)
        if current_section_text:
            chunks.append(Document_Chunk(
                index=len(chunks), 
                page=section_index + 1, 
                chunk_type=Chunk_Type.PARAGRAPH, 
                text=current_section_text.strip(),
                tokens_est=current_section_tokens
            ))

        
        for table_num, table in enumerate(doc.tables):
            table_text = "\n".join([" | ".join([cell.text.strip() for cell in row.cells]) for row in table.rows])
            if table_text.strip():
                chunks.append(Document_Chunk(
                    index=len(chunks),
                    page=table_num + 1,
                    chunk_type = Chunk_Type.TABLE,
                    text=table_text,
                    tokens_est=int(len(table_text.split())*1.3)
                ))
            
            # For figures: extract caption only for now (fast & useful)
            # (PDF: fitz can pull image blocks; DOCX: inline_shapes captions)
            # Example stub:
            # chunks.append(Document_Chunk(..., chunk_type=Chunk_Type.FIGURE, text=caption))
    else: 
        raise ValueError(f"Unsupported file type: {suffix}")
    
    return chunks